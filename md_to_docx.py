from pathlib import Path
from docx import Document

src = Path('AUDIT_REPORT.md')
dst = Path('AUDIT_REPORT.docx')
text = src.read_text(encoding='utf-8')
lines = text.splitlines()
doc = Document()
paragraph = None
in_code = False
for line in lines:
    stripped = line.strip()
    if stripped.startswith('```'):
        in_code = not in_code
        if in_code:
            paragraph = doc.add_paragraph()
            paragraph.style = 'Intense Quote'
        continue
    if in_code:
        paragraph.add_run(line + '\n')
        continue
    if stripped == '':
        paragraph = None
        continue
    if stripped.startswith('# '):
        paragraph = doc.add_paragraph(stripped[2:].strip(), style='Heading 1')
        continue
    if stripped.startswith('## '):
        paragraph = doc.add_paragraph(stripped[3:].strip(), style='Heading 2')
        continue
    if stripped.startswith('### '):
        paragraph = doc.add_paragraph(stripped[4:].strip(), style='Heading 3')
        continue
    if stripped.startswith('- '):
        paragraph = doc.add_paragraph(stripped[2:].strip(), style='List Bullet')
        continue
    if stripped.startswith('* '):
        paragraph = doc.add_paragraph(stripped[2:].strip(), style='List Bullet')
        continue
    if len(stripped) >= 4 and stripped[:2].isdigit() and stripped[2:4] == '. ':
        paragraph = doc.add_paragraph(stripped[4:].strip(), style='List Number')
        continue
    if stripped.startswith('> '):
        paragraph = doc.add_paragraph(stripped[2:].strip(), style='Intense Quote')
        continue
    paragraph = doc.add_paragraph(line)

try:
    doc.save(dst)
    print('OK')
except Exception as e:
    print('ERROR', e)
